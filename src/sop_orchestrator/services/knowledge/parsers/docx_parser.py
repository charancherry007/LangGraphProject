import docx
import hashlib
from pathlib import Path
from datetime import datetime
from src.sop_orchestrator.models.knowledge import ParsedDocument
from src.sop_orchestrator.services.knowledge.parsers.base import BaseParser
from src.sop_orchestrator.services.knowledge.exceptions import DocumentParsingException

class DOCXParser(BaseParser):
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path, document_id: str, relative_path: str) -> ParsedDocument:
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            
            normalized_text = " ".join(raw_text.split())
            hash_value = hashlib.sha256(raw_text.encode("utf-8")).hexdigest()
            
            stat = file_path.stat()
            
            return ParsedDocument(
                document_id=document_id,
                document_name=file_path.name,
                document_type="DOCX",
                relative_path=relative_path,
                created_at=datetime.fromtimestamp(stat.st_ctime),
                modified_at=datetime.fromtimestamp(stat.st_mtime),
                page_count=None,
                raw_text=raw_text,
                normalized_text=normalized_text,
                hash_value=hash_value,
                paragraphs=paragraphs
            )
        except Exception as e:
            raise DocumentParsingException(f"Failed to parse DOCX {file_path}: {e}")
